from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm , Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml.etree import QName
import re
import comtypes.client
import os
import time
#from devCode import FormatTableStyle, RemoveTable , RemoveSectionBreaks, CheckSmallCapsFunction
#from devCode import replace_bullets 
from logs.logs_handler import get_logger
import common_func.word_paragraph_operations as word_para_ops
import common_func.word_bullets_operations as word_bull_ops
import common_func.word_table_operations as word_table_ops
import common_func.word_effects_CheckSmallCapsFunction as word_effects_checkSmallCaps

logger = get_logger(__name__)


def format_document(doc_path,region_code = None):
    """
    Load a DOCX document, apply formatting and modifications, and save it.
    """
    # Load the document
    doc = Document(doc_path)

    # Apply formatting to all paragraphs
    for paragraph in doc.paragraphs:
        word_para_ops.format_paragraph(paragraph)
        word_para_ops.align_paragraph(paragraph)
        word_bull_ops.apply_bullet_formatting(paragraph)

    # Apply formatting to all tables
    for table in doc.tables:
        word_table_ops.indent_table_left(table)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    word_para_ops.format_paragraph(paragraph)
                    word_para_ops.align_paragraph(paragraph)
                    word_bull_ops.apply_bullet_formatting(paragraph)

    word_para_ops.remove_spacing_between_paras(doc)
    # remove_tabs_from_document(doc)
    
    # Save the modified document
    doc.save(doc_path)
    logger.info(f" Formatted document and saved: {doc_path}")

    # custom call functions
    word_table_ops.add_border_doc(doc_path,doc_path)        #adds border to all the table elements
    word_table_ops.replace_image_with_text(doc_path,"{Non Displayable Image}",doc_path)       # replaces image with "{Non Displayable Image}"
    word_table_ops.remove_line_numbering(doc_path, doc_path)   
    word_bull_ops.process_bullets_in_document(doc_path)   
    word_bull_ops.set_bullet_follow_char_to_space(doc_path) 
    word_bull_ops.set_bullet_font(doc_path)    
    #remove_header_footer_sections(doc_path,region_code)
    word_effects_checkSmallCaps.check_small_caps(doc_path)
    word_para_ops.remove_section_breaks(doc_path)
    word_para_ops.word_file_indentation(doc_path)                             
    

# def remove_header_footer_sections(doc_path,region_code):
#     # Open the document
#     doc = Document(doc_path)

#     # Remove headers and footers
#     for section in doc.sections:
#         # Remove header references
#         section.different_first_page_header_footer = False
#         section.header.is_linked_to_previous = True
        
#         # Remove footer references
#         section.footer.is_linked_to_previous = True

#         # Clear any existing header and footer content
#         section.header.paragraphs.clear()
#         if region_code is not None:
#             if region_code.strip().lower() != 'ca':        
#                 section.footer.paragraphs.clear()

#     # Remove header and footer references
#     for section in doc.sections:
#         sectPr = section._sectPr
#         for child in sectPr:
#             if child.tag.endswith(('headerReference', 'footerReference')):
#                 sectPr.remove(child)

#     # Set header and footer distance to 0
#     for section in doc.sections:
#         section.header_distance = 0
#         section.footer_distance = 0

#     # Save the modified document
#     doc.save(doc_path)
#     logging.info(f"Header and footer sections removed from the document: {doc_path}")


def clean_up_temp_file(temp_file_path):
    """
    Remove the temporary DOCX file if it exists.
    """
    if os.path.exists(temp_file_path):
        os.remove(temp_file_path)

# def remove_tabs_from_document(doc):
#     """Function to remove tabs from the file text while preserving images"""

#     def process_paragraph(paragraph):
#         for run in paragraph.runs:
#             if not run.element.findall('.//w:drawing', namespaces=run.element.nsmap):
#                 # Only process runs that don't contain images
#                 run.text = run.text.replace('\t', ' ')

#     # Process paragraphs
#     for paragraph in doc.paragraphs:
#         process_paragraph(paragraph)

#     # Process tables
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     process_paragraph(paragraph)

#     logging.info(f"Tabs removed successfully")


# def check_file_type_and_convert(input_file_path: str, output_doc_path: str):
#     """
#     Check the file type and convert it to DOCX format, then apply formatting.
#     """
#     if os.path.exists(output_doc_path):
#         os.remove(output_doc_path)

#     base_name, file_extension = os.path.splitext(input_file_path)
#     temp_doc_path = f"{base_name}.docx"

#     if file_extension.strip().lower() in ['.doc', '.pdf', '.docx']:
#         convert_file_to_docx(input_file_path, temp_doc_path)
#     else:
#         logger.warning("Unsupported file type. Only .doc and .pdf are supported.")
#         return None

#     time.sleep(5)  # Ensure the file is fully converted before proceeding
#     return temp_doc_path

# format_document(r"C:\File\NETSCAN\Input\NETSCAN_CO_Test_14\Exception\2025-00021_co_p.docx")