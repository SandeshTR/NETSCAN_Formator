from bs4 import BeautifulSoup
from docx import Document 
#from HTML2Docx import html2docx 
import os


def HTMLtoWord(html_file_path,word_file_path):

    HTMLParsing(html_file_path,html_file_path.replace(".html","_1.html"))

    html_file_path=html_file_path.replace(".html","_1.html")
     #Read the HTML content from the file
    with open(html_file_path, 'r',encoding='utf-8') as file:
        html_content = file.read()
        print(f"html content is \n : {html_content}")

    # Parse the HTML content with BeautifulSoup
    soup = BeautifulSoup(html_content, 'lxml')
    print(f"content is \n : {soup}")
    # Create a new Word document
    doc = Document()
    #doc.add_heading('Converted HTML', level=1)

    # Function to handle different HTML elements
    def add_element_to_doc(element):
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'p' :
            data=doc.add_paragraph()
            run=data.add_run(element.get_text())
            run.bold = True            
        elif element.name == 'hr':
            data=doc.add_paragraph()
            run=data.add_run(element.get_text())
            run.bold = True
        elif element.name == 'table':
            # Handle tables
            rows = element.find_all('tr')
            if rows:
                flagnext=False
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                for row_index, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for col_index, cell in enumerate(cells):
                        #table.cell(row_index, col_index).text = cell.get_text(strip=True)
                        
                        cell_text = cell.get_text(strip=True)
                        if flagnext==True:
                            flagnext=False
                            #if 0 <= row_index < len(table.rows):
                                #table._tbl.remove(table.rows[row_index]._tr)
                            continue

                        if str(cell_text).lower()=="rule" or str(cell_text).lower()=="adopted rules":
                            flagnext=True
                            continue
                        
                        table.style='Table Grid'
                        if (str(cell_text).lower()=="additional information" or str(cell_text).lower()=="basis and purpose" or 
                        str(cell_text).lower()=="redline" or str(cell_text).lower()=="emergency justification") : 
                            cell_text=cell_text+"(see below if available)"

                        cell_paragraph = table.cell(row_index, col_index).paragraphs[0]
                        run = cell_paragraph.add_run(cell_text)

                        if cell.find('strong') or cell.find('b'):                            
                            run.bold = True   

                        if cell.find('i') or cell.find('em'):                            
                            run.italic=True

                        #if cell.find('u'):                            
                        #    run.underline=True
                                             
                                               
                        

        elif element.name == 'img':
            # Handle images
            img_url = element.get('src')
            if img_url:
                doc.add_paragraph('Image: {}'.format(img_url))
        # Add more cases as needed for other elements

    # Iterate over all elements in the HTML
    for element in soup.body.find_all(True):  # Find all tags within body
        add_element_to_doc(element)

    # Save the Word document
    doc.save(word_file_path)
    os.remove(html_file_path)

    print(f'HTML content has been converted and saved to {word_file_path}')

def HTMLParsing(htmlpath,output_html_file_path):
    # Replace with the path to your local HTML file
    input_html_file_path =  htmlpath
    # Replace with the path where you want to save your output HTML file
    #output_html_file_path = docxpath

    # Read the HTML content from the input file
    with open(input_html_file_path, 'r',encoding='utf-8') as file:
        html_content = file.read()

    # Parse the HTML content with BeautifulSoup
    soup = BeautifulSoup(html_content, 'lxml')
    print(f"content is 1 soup \n : {soup}")

    # Find all <table> elements
    tables = soup.find_all('table')
    print(f"tables are : \n{tables}")
    # Create a new BeautifulSoup object for the output HTML
    output_soup = BeautifulSoup('<html><head><title>Extracted Tables</title></head><body></body></html>', 'lxml')
    output_body = output_soup.body
    extracted_texts = []
    # Add each table to the new HTML
    for table_index, table in enumerate(tables):
        # Add a heading for each table
        if table_index<2:
            continue
        
        #heading = output_soup.new_tag('hr', style='border: 3px solid black; margin: 20px 0;')
        #heading = output_soup.new_tag('p')
        #heading= output_soup.new_tag('strong')
        if table_index == 2:              
            rows = table.find_all('tr')
            for row in rows:
                # Find all cells in the row
                cells = row.find_all('td')
                #print(str(cells))
                # Loop through each cell
                for cell in cells:                
                    # Check if the cell contains the text "Details of Tracking Number"
                    if 'Details of Tracking Number' in cell.get_text():                       
                        data = str(cell.get_text())
                        start_marker = 'Details of Tracking Number'
                        end_marker = 'CCR details' 
                        start_index = data.find(start_marker)
                        end_index = data.find(end_marker, start_index)

                        soup1 = BeautifulSoup(str(table), 'html.parser')
                        #print(str(cell.get_text()))
                        # Initialize an array to store the extracted text
                        

                        # Find all <span> elements with the class 'darkBlueText'
                        span_tags = soup1.find_all('span', class_='darkBlueText')

                        for span in span_tags:
                            # Extract text from each <span> element
                            text = span.get_text(strip=True)
                            extracted_texts.append(text)

                        print("Extracted texts:", extracted_texts[0]) 

                        if start_index != -1 and end_index != -1:
                            extracted_text = data[start_index:end_index].strip()                               
                            heading1 = output_soup.new_tag('p')
                            heading1.string= extracted_text.replace('\n\t\xa0','')
                            output_body.append(heading1) 
                        
                    break
        elif table_index>2:
            heading = output_soup.new_tag('p')
            if table_index == 3:
                heading.string =extracted_texts[0]

            elif table_index == 4:
                heading.string =extracted_texts[1]

            elif table_index == 5:
                heading.string =extracted_texts[2]

            elif table_index == 6:
                heading.string =extracted_texts[3]
            else:
                print('Zero')

        #heading.string = f"Table {table_index + 1}"
            output_body.append(heading)          

            # Append the table
            output_body.append(table)

    # Write the output HTML to a file
    with open(output_html_file_path, 'w',encoding='utf-8') as file:
        file.write(str(output_soup.prettify()))

     


#print(f'Tables have been extracted and saved to {docxpath}')
#HTMLtoWord(docxpath,docxpath.replace(".html",".doc"))
#os.remove(docxpath)

#docxpath=r"C:\File\Image\Samples\CO Sample files- Prog conversion\NETSCAN_CO.netscanregs.071524.000026-input\NETSCAN_CO240712cas\co_a005a207.html"
#HTMLtoWord(docxpath,docxpath.replace(".html",".docx"))
