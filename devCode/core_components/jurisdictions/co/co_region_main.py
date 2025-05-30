from docx import Document
from docx.oxml.ns import qn
import pytesseract
from PIL import Image, ImageFilter
import io, os
from docx.oxml import OxmlElement
import re
from core_components.jurisdictions.co import co_redline
from core_components.jurisdictions.co import co_aft
import logging

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(filename)s - %(message)s')

# def replace_bullets_with_text(input_docx_path):
#     "This function is used to check bullets in the docx file and replace them with value "
#     doc = Document(input_docx_path)
#     logging.info('replaced bullet with text, normal way')
#     bullet_characters = ['•', '◦', '▪', '▫','●','►', '❖', '❥', '➤', '➢', '⦿']
    
#     def replace_bullet_in_paragraph(para):
#         for bullet in bullet_characters:
#             if para.text.strip().startswith(bullet):
#                 for run in para.runs:
#                     if run.text.startswith(bullet):
#                         run.text = run.text.replace(bullet, '~#8226;', 1)
#                         break
    
#     for para in doc.paragraphs:
#         replace_bullet_in_paragraph(para)
    
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for para in cell.paragraphs:
#                     replace_bullet_in_paragraph(para)
    
#     doc.save(input_docx_path)
#     logging.info(f"Bullet module document saved as {input_docx_path}")


def remove_pagechars(input_docx_path):
    """Remove certain extra characters in file"""
    doc = Document(input_docx_path)
    pattern = re.compile(r'\b(SBP \d+|RA \d+|SE \d+)\b')
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = pattern.sub('',run.text)
            # Remove leading special characters
            run.text = re.sub(r'^[^\w\s]+', '', run.text)
            # Remove leading numbers
            run.text = re.sub(r'^\d+\s*', '', run.text)
    doc.save(input_docx_path)      
    logging.info(f'removed page characters and saved as {input_docx_path} ')

def convert_hyperlink_to_text(input_docx_path):
    doc = Document(input_docx_path)
    for paragraph in doc.paragraphs:
        p = paragraph._p
        hyperlinks = p.xpath(".//w:hyperlink")
        for hyperlink in hyperlinks:
            # Check if the hyperlink is a bookmark or footnote reference
            if hyperlink.get(qn('w:anchor')) is not None:
                logging.info("Skipping hyperlink that references the same document.")
                continue
            
            # Get the text in the hyperlink
            hyperlink_text = ''.join([node.text for node in hyperlink.xpath('.//w:t')])
            logging.info(hyperlink_text)
            
            # Create a new run with the hyperlink text
            new_run = OxmlElement('w:r')
            new_text = OxmlElement('w:t')
            new_text.text = hyperlink_text
            new_run.append(new_text)

            # Create the rPr element to hold the formatting properties
            rPr = OxmlElement('w:rPr')
            
            # Set the font to Times New Roman
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(rFonts)
            
            # Set the font size to 12 (which is 24 half-points)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '24')
            rPr.append(sz)
            
            # Append the rPr element to the run
            new_run.insert(0, rPr)

            # Insert the new run before the hyperlink
            hyperlink.addprevious(new_run)
            
            # Remove the hyperlink
            p.remove(hyperlink)
            logging.info("Hyperlink converted to text.")
    
    doc.save(input_docx_path)
    logging.info(f"Removed hyperlinks and saved as {input_docx_path}")


def convert_hyperlink_to_text_old(input_docx_path):
    '''Deprecation Warning :This function is used to convert hyperlinks to text in the docx file
    '''
    doc = Document(input_docx_path)
    for paragraph in doc.paragraphs:
        p = paragraph._p
        hyperlinks = p.xpath(".//w:hyperlink")
        for hyperlink in hyperlinks:
            # Get the text in the hyperlink
            hyperlink_text = ''.join([node.text for node in hyperlink.xpath('.//w:t')])
            print(hyperlink_text)
            
            # Create a new run with the hyperlink text
            new_run = OxmlElement('w:r')
            new_text = OxmlElement('w:t')
            new_text.text = hyperlink_text
            new_run.append(new_text)

            # Create the rPr element to hold the formatting properties
            rPr = OxmlElement('w:rPr')
            
            # Set the font to Times New Roman
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(rFonts)
            
            # Set the font size to 12 (which is 24 half-points)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '24')
            rPr.append(sz)
            
            # Append the rPr element to the run
            new_run.insert(0, rPr)

            # Insert the new run before the hyperlink
            hyperlink.addprevious(new_run)
            
            # Remove the hyperlink
            p.remove(hyperlink)
            print('this ran')
    doc.save(input_docx_path)
    logging.info(f"remove hyperlinks and saved as {input_docx_path}")



# def remove_tabs_from_document(input_docx_path):
#     """Function to remove tabs from the file text while preserving images"""
#     doc = Document(input_docx_path)

#     def process_paragraph(paragraph):
#         for run in paragraph.runs:
#             if not run.element.findall('.//w:drawing', namespaces=run.element.nsmap):
#                 # Only process runs that don't contain images
#                 run.text = run.text.replace('\t', '')

#     # Process paragraphs
#     for paragraph in doc.paragraphs:
#         process_paragraph(paragraph)

#     # Process tables
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     process_paragraph(paragraph)

#     doc.save(input_docx_path)
#     logging.info(f"Tabs removed and document saved at {input_docx_path}")


def extract_images_with_locations(input_docx_path,header_check=False):
    logging.info(f"header check : {header_check}")
 
    doc = Document(input_docx_path)
    images_with_locations = []
    index = 0
    logging.info('read document')
    doc_part_to_consider = doc.paragraphs[:11] if header_check else doc.paragraphs  #depending on weather you want to check header or all the values use this
     

    for i, paragraph in enumerate(doc_part_to_consider):
        for element in paragraph._element.iter():
            if element.tag == qn('w:drawing'):
                blip = element.find('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                if blip is not None:
                    img_id = blip.get(qn('r:embed'))
                    image_part = doc.part.related_parts[img_id]
                    images_with_locations.append((index, i, image_part.blob))
                    index += 1

    logging.info(f'number of images identified {len(images_with_locations)}')
    return images_with_locations



def extract_text_from_images(images_with_locations):
    """extract images and there location from docx """
    texts = []
    for idx, location, img_blob in images_with_locations:
        image = Image.open(io.BytesIO(img_blob))
        text = pytesseract.image_to_string(image=image)
        print(f'text extracted is : {text}')
        if text.strip():                       #filtering and adding only non empty image locations
            texts.append((idx, location, text))
            logging.info(f'added location text {location}')
    return texts


def insert_single_text_in_paragraph(input_docx_path, location, text):
    """This function adds logo text present at the top of page along with 
    file type for ADDITIONAL INFORMATION/ EMERGENCY JUSTIFICATION/ BASIS AND PURPOSE files"""
    
    base_name, file_extension = os.path.splitext(input_docx_path)
    doc = Document(input_docx_path)
                                 #only one logo present in docx file
    paragraph = doc.paragraphs[location]
    paragraph.clear()
    logging.info(" cleared text - removed logo content ")
    for run in paragraph.runs:
        if '_' in run.text:
            run.text = run.text.replace('_', '')
    if location < 3:
        if 'addinfo' in base_name.strip().lower():
            paragraph.add_run('ADDITIONAL INFORMATION \n' + text.strip().replace('\n\n','\n')).bold = True
        elif 'emergency' in base_name.strip().lower():
            paragraph.add_run('EMERGENCY JUSTIFICATION \n' + text.strip().replace('\n\n','\n')).bold = True    
        elif 'basisandpurpose' in base_name.strip().lower():
            paragraph.add_run('BASIS AND PURPOSE\n'+ text.strip().replace('\n\n','\n')).bold = True         
        else:
            paragraph.add_run(text.strip().replace('\n\n','\n')).bold = True

    doc.save(input_docx_path)
    logging.info(f"Modified document saved as {input_docx_path}")




def insert_multiple_text_in_paragraph(input_docx_path,img_location_list):
    """Function to add multiple text based on the list index position provided"""
    doc = Document(input_docx_path)
    
    sorted_list = sorted(img_location_list,key=lambda x: x[1],reverse=True)
    logging.info(f'sorted list {sorted_list}')

    for index,location,content in sorted_list:
        if "colorado" in content.lower():
            paragraph = doc.paragraphs[location]

            # Remove all runs (including images) from the paragraph
            for run in paragraph.runs:
                run.clear()
            
            paragraph.add_run(content.strip().replace('\n\n', '\n')).bold = True
            logging.info(f'Added colorado data {content}')
        else:
            logging.info(f"Skipping content at location {location} as it doesn't contain 'Colorado'")
    
    base_name = os.path.basename(input_docx_path).strip().lower()

    if any(name in base_name.strip().lower() for name in ['addinfo','emergency','basisandpurpose']):
        paragraph = doc.paragraphs[0].insert_paragraph_before()
        if 'addinfo' in base_name.strip().lower():
            paragraph.add_run('ADDITIONAL INFORMATION').bold = True
        elif 'emergency' in base_name.strip().lower():
            paragraph.add_run('EMERGENCY JUSTIFICATION').bold = True
        elif 'basisandpurpose' in base_name.strip().lower():
            paragraph.add_run('BASIS AND PURPOSE').bold = True

    
    # Save the modified document
    doc.save(input_docx_path)      


def add_first_line_header(input_docx_path):
    base_name, file_extension = os.path.splitext(input_docx_path)

    doc = Document(input_docx_path)   
    
    paragraph = doc.paragraphs[0].insert_paragraph_before()

    if 'addinfo' in base_name.strip().lower():
        if doc.paragraphs[0].text =='ADDITIONAL INFORMATION' or doc.paragraphs[1].text == 'ADDITIONAL INFORMATION' or doc.paragraphs[2].text == 'ADDITIONAL INFORMATION':
            return
        paragraph.add_run('ADDITIONAL INFORMATION').bold = True
    elif 'emergency' in base_name.strip().lower():
        paragraph.add_run('EMERGENCY JUSTIFICATION').bold = True    
    elif 'basisandpurpose' in base_name.strip().lower():
        paragraph.add_run('BASIS AND PURPOSE').bold = True         
    doc.save(input_docx_path)
    logging.info(f"Added first line info when images count is zero ,document saved as {input_docx_path}")



def main_co_files(input_docx_path,source_file_path):
    """Main function for co region, this function assigns and calls functions based """
    
    logging.info(f"Starting main_co_files function with input: {input_docx_path}")
    text_inserted = False
    #get basename of docx
    base_name, file_extension = os.path.splitext(input_docx_path)

    #logic to check initial header/logo present and add text at that location
    images_with_locations = extract_images_with_locations(input_docx_path,header_check=False)
    
    logging.info(f"Number of images extracted: {len(images_with_locations)}")

    if len(images_with_locations) > 0:  # File contains headers
        img_texts = extract_text_from_images(images_with_locations)
        texts = [text for text in img_texts  if text[2]]
        logging.debug(f"Extracted texts: {texts}")


        if len(texts)>0:
            if len(images_with_locations) > 1 and len(texts) > 1:
                # This is for files which has logo to begin with
                logging.info("calling multiple text insert")
                if any('colorado' in text[2].strip().lower() for text in texts):
                    insert_multiple_text_in_paragraph(input_docx_path=input_docx_path,img_location_list=img_texts )
            else:
                logging.info("calling single text insert")
                #need to check what basename is aft or kinda 
                insert_single_text_in_paragraph(input_docx_path=input_docx_path,location=texts[0][1],text=texts[0][2])                
                text_inserted = True

    #check file type and perform action
    if 'aft' in base_name.strip().lower():
        #chek pdf present as source and if keywords present
        co_aft.determine_aft_file_type(input_docx_path=input_docx_path,source_file_path=source_file_path)

    elif 'redline' in base_name.strip().lower():
        #calling redline code
        co_redline.main(input_docx_path) 

    elif 'addinfo' in base_name.strip().lower():   #Added extra check for addinfo as no data was getting added (added on 2025-01-16)
        if not text_inserted:
            # add_first_line_header(input_docx_path)
            if len(images_with_locations) > 0 and len(texts) > 0:
                insert_single_text_in_paragraph(input_docx_path=input_docx_path,location=0,text=texts[0][2])

    if len(images_with_locations) <= 0:
        add_first_line_header(input_docx_path=input_docx_path)

    # # run common guideline logics
    # replace_bullets_with_text(input_docx_path)
    convert_hyperlink_to_text(input_docx_path)     # Issue Cause is this function which is not working properly for linked footers
    # remove_tabs_from_document(input_docx_path)