from docx import Document
from fuzzywuzzy import fuzz 
from logs.logs_handler import get_logger

logger = get_logger(__name__)

def find_duplicate_sections(doc, doc_type):
    """
    Find potential duplicate sections in the document based on the document type.
    Returns:
    dict: A dictionary containing the text and index ranges of potential duplicate sections.
    """
    duplicate_info = {'text': [], 'index': []}
    end_of_text_state_basis_index = []

    # Determine the end text lookup based on document type
    end_text_lookup = 'secretary of state' if doc_type == 'TYPE 1' else 'statement of basis and purpose'
    
    # Find all occurrences of the end text
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        if end_text_lookup in text:
            end_of_text_state_basis_index.append(i)
    
    # For each end text occurrence, find the corresponding start ("title of rule")
    for n in end_of_text_state_basis_index:
        for i in range(n-1, -1, -1):
            text = doc.paragraphs[i].text.strip().lower()
            if 'title of rule' in text:
                duplicate_info['index'].append([i, n-1])
                break

        # Extract the text between start and end
        search_text = " ".join([p.text.strip().replace('\t', '') for p in doc.paragraphs[i:n]]).strip()
        duplicate_info['text'].append(search_text)
    return duplicate_info


def find_duplicate_text(doc, start_index, search_text):
    """
    Find duplicate text in the document starting from a given index.
    Returns:
    dict: A dictionary containing the text and index ranges of found duplicates.
    """
    duplicate_info = {'text': [], 'index': []}
    for i, para in enumerate(doc.paragraphs[start_index:], start=start_index):
        text = para.text.strip().lower()

        # Mark the start of a potential duplicate section
        if 'title of rule' in text:
            if duplicate_info['index']:
                duplicate_info['index'].pop()
            duplicate_info['index'].append(i)

        # Check for the end of a potential duplicate section
        if 'division / contact / phone:' in text:
            temp_index = i
            match_text = [p.text.strip().replace('\t', '') for p in doc.paragraphs[duplicate_info['index'][0]:temp_index]]
            
            # Check the next few paragraphs to ensure we capture the full duplicate section
            for no, p in enumerate(doc.paragraphs[temp_index:temp_index+6], start=temp_index):
                match_text.append(p.text.replace('\t', ''))
                str_match_text = ''.join(match_text)
                
                # Use fuzzy matching to identify duplicates
                if fuzz.ratio(search_text, str_match_text) > 95:
                    duplicate_info['index'].append(no)
                    break

            if len(duplicate_info['index']) > 1:
                break
            if i == len(doc.paragraphs) - 1:
                return duplicate_info
            
    # If a duplicate is found, extract its text
    if len(duplicate_info['index']) > 1:
        start = duplicate_info['index'][0]
        end = duplicate_info['index'][1] + 1
        text = ''.join([doc.paragraphs[row_no].text.strip().replace('\t', '') for row_no in range(start, end)])
        duplicate_info['text'].append(text)
    return duplicate_info



def remove_paragraphs(doc, indexes):
    """
    Remove paragraphs from the document based on the provided indexes.
    """
    # Sort indexes in reverse order to avoid shifting issues when removing paragraphs
    sorted_indexes = sorted(indexes, key=lambda x: x[0], reverse=True)
    for index in sorted_indexes:
        print(index)
        if len(index) > 1:
            for i in range(index[1], index[0]-1, -1):
                p = doc.paragraphs[i]._element
                p.getparent().remove(p)
                p._element = p._p = None