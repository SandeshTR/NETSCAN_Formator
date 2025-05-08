from docx import Document
import os
import pdfplumber
import logging
from logs.logs_handler import get_logger
from core_components.jurisdictions.co.co_common_func import find_duplicate_sections, find_duplicate_text , remove_paragraphs

logger = get_logger(__name__)


def extract_text_from_pdf(input_file):
    """
    Extract text from a PDF file and return the text up to the first non-empty line after an empty line in the half-page of the first page.
    """
    # Get the basename and extension of the input file 
    basename, extension = os.path.splitext(input_file)

    # Check if 'aft' is in the basename and the file is a PDF
    if 'aft' in basename.lower().strip() and extension.lower() == '.pdf':
        with pdfplumber.open(input_file) as pdf:
            page_1 = pdf.pages[0]
            half_page = page_1.crop((0, 0, page_1.width, page_1.height / 2 + 100))
            text = half_page.extract_text(layout = True)
            standard_text = half_page.extract_text()

            # Split the text by 'Subject:'
            lines = text.split('Subject:')[1].split('\n')
            print(lines)

            # Initialize variables to track state
            found_empty_line = False
            non_empty_line_after_empty = None

            # Iterate through lines to find the first non-empty line after an empty line
            for line in lines:
                if found_empty_line:
                    if line.strip():                                           # Check if the line is not empty or just whitespace
                        non_empty_line_after_empty = line.strip()
                        break
                else:
                    if not line.strip():                                        # Check if the line is empty or just whitespace
                        found_empty_line = True

            # Return the text up to the first non-empty line after an empty line
            if non_empty_line_after_empty:
                print(f'partly extracted text : {standard_text.split(non_empty_line_after_empty.strip())[0]}')
                return standard_text.split(non_empty_line_after_empty.strip())[0]
            else:
                return ""
            

def remove_duplicate_paragraphs(docx_path, doc_type):
    """
    Remove duplicate paragraphs from the document.
    """
    doc = Document(docx_path)

    # Find potential duplicate sections
    master_dict_text_index = find_duplicate_sections(doc, doc_type)
    print(f'master_dict_text_index {master_dict_text_index}')

    del_dup_index_list = []

    # For each potential duplicate section, search for actual duplicates
    for index, (start, end) in enumerate(master_dict_text_index['index']):
        current_start = end + 1
        while True:
            duplicate_info_dic = find_duplicate_text(doc, start_index=current_start, search_text=master_dict_text_index['text'][index])
            if len(duplicate_info_dic['index']) > 1:
                del_dup_index_list.append(duplicate_info_dic['index'])
                current_start = duplicate_info_dic['index'][-1] + 1
            else:
                break
    
    # Remove the identified duplicate paragraphs
    remove_paragraphs(doc, del_dup_index_list)
    doc.save(docx_path)




def check_text_present(values:list,text:str):
    return all(substring in text for substring in values)



def determine_aft_file_type(input_docx_path,source_file_path):
    """
    Determine the type of AFT file and process it accordingly.
    """

    #Varient 2!
    #dealing with title of rule and rule number varients
    doc = Document(input_docx_path)

    # This flow is for the files that has no logo and has title and rule number files
    doc_type = ''
    index_of_title = 0
    
    # Find the index of "Title of Rule"
    for cntr, para in enumerate(doc.paragraphs):
        if 'title of rule:' in para.text.strip().lower():
            index_of_title = cntr
            break
    
    # Check the next few paragraphs for key phrases
    text_check = []
    for par in doc.paragraphs[index_of_title:index_of_title+10]:
        text_check.append(par.text)

    del doc   #discarding memory data    
    check_text_val = " ".join(text_check).strip().lower()

    title_and_rule_number = 'title of rule:' in check_text_val and 'rule number:' in check_text_val

    # Determine the document type based on key phrases
    if title_and_rule_number and 'division / contact / phone:' in check_text_val and 'secretary of state' in check_text_val:
        doc_type = 'TYPE 1'
    elif title_and_rule_number and 'division / contact / phone:' in check_text_val and 'statement of basis and purpose' in check_text_val:
        doc_type = 'TYPE 2'
    elif title_and_rule_number and 'Compliance and Innovation Division (CID) / Matt Bohanan /'.strip().lower() in check_text_val and 'statement of basis and purpose' in check_text_val:
        doc_type = 'TYPE 3'
    elif title_and_rule_number and 'OCL / Tiffani Domokos and Cassandra Keller'.strip().lower() in check_text_val and 'statement of basis and purpose' in check_text_val:
        doc_type = 'TYPE 4'
    elif 'Add any new varients here':
            pass
    
    logging.info(f'type assigned : {doc_type}')
    match doc_type:
        case 'TYPE 1' | 'TYPE 2':
            remove_duplicate_paragraphs(input_docx_path, doc_type)
        case 'TYPE 3':
            pass
        case _:
            logging.info('no matching varient')

    

#file_path = r"C:\Users\6120867\Downloads\Netscan_Test_Files\co_files\NETSCAN_CO240701cas_Input\co_p002aft183.pdf"
# file_path = r"C:\Users\6120867\Downloads\Netscan_Test_Files\co_files\NETSCAN_CO240701cas_Input\co_p002aft183.docx"
# file_path = r"C:\Users\6120867\Downloads\Netscan_Test_Files\co_files\NETSCAN_CO240701cas_Input\co_p001aft183.docx"
# determine_aft_file_type(file_path,file_path.replace('.docx','.pdf'))