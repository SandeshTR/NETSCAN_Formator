from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm
from fuzzywuzzy import fuzz
import logging



# def get_first_header_text(doc):
#     """ Get the text of the first header """
#     if doc.sections[0].header.is_linked_to_previous:
#         header = doc.sections[0].header.paragraphs
#     else:
#         header = doc.sections[0].header.paragraphs
#     logging.info(f"header is {header}")
#     if header:
#         return header[0].text
#     return None

# def remove_headers_and_footers(doc):
#     """ remove all headers and footers in the document """
#     namespace = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
#     headerReference_tag = qn('w:headerReference')
#     footerReference_tag = qn('w:footerReference')

#     for section in doc.sections:
#         sectPr = section._sectPr

#         # Remove header references
#         header_refs = sectPr.findall(headerReference_tag)
#         for header_ref in header_refs:
#             sectPr.remove(header_ref)

#          # Remove footer references
#         footer_refs = sectPr.findall(footerReference_tag)
#         for footer_ref in footer_refs:
#             sectPr.remove(footer_ref)

def add_text(doc, first_header_text=''):
    """Insert a new paragraph at the beginning of the document"""
    if first_header_text:
        new_paragraph = doc.paragraphs[0].insert_paragraph_before(f'REDLINE \n{first_header_text}')
    else:
        new_paragraph = doc.paragraphs[0].insert_paragraph_before('REDLINE')
    logging.info("'REDLINE' text added to document")
    
    # Make the text bold
    for run in new_paragraph.runs:
        run.bold = True

    new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    new_paragraph.paragraph_format.space_before = Pt(3)
    new_paragraph.paragraph_format.space_after = Pt(3)
    new_paragraph.paragraph_format.left_indent = Cm(0)
    new_paragraph.paragraph_format.first_line_indent = Cm(0)
    logging.info('Made text, para alignment')




def find_duplicate_sections(doc, doc_type):
    """
    Find potential duplicate sections in the document based on the document type.
    Returns:
    dict: A dictionary containing the text and index ranges of potential duplicate sections.
    """
    duplicate_info = {'text': [], 'index': []}
    end_of_text_state_basis_index = []

    # Determine the end text lookup based on document type
    end_text_lookup = 'secretary of state' if doc_type == 'TYPE 1' else 'statement of basis and purpose'
    
    # Find all occurrences of the end text
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        if end_text_lookup in text:
            end_of_text_state_basis_index.append(i)
    
    # For each end text occurrence, find the corresponding start ("title of rule")
    for n in end_of_text_state_basis_index:
        for i in range(n-1, -1, -1):
            text = doc.paragraphs[i].text.strip().lower()
            if 'title of rule' in text:
                duplicate_info['index'].append([i, n-1])
                break

        # Extract the text between start and end
        search_text = " ".join([p.text.strip().replace('\t', '') for p in doc.paragraphs[i:n]]).strip()
        duplicate_info['text'].append(search_text)
    return duplicate_info



def find_duplicate_text(doc, start_index, search_text):
    """
    Find duplicate text in the document starting from a given index.
    Returns:
    dict: A dictionary containing the text and index ranges of found duplicates.
    """
    duplicate_info = {'text': [], 'index': []}
    for i, para in enumerate(doc.paragraphs[start_index:], start=start_index):
        text = para.text.strip().lower()

        # Mark the start of a potential duplicate section
        if 'title of rule' in text:
            if duplicate_info['index']:
                duplicate_info['index'].pop()
            duplicate_info['index'].append(i)

        # Check for the end of a potential duplicate section
        if 'division / contact / phone:' in text:
            temp_index = i
            match_text = [p.text.strip().replace('\t', '') for p in doc.paragraphs[duplicate_info['index'][0]:temp_index]]
            
            # Check the next few paragraphs to ensure we capture the full duplicate section
            for no, p in enumerate(doc.paragraphs[temp_index:temp_index+6], start=temp_index):
                match_text.append(p.text.replace('\t', ''))
                str_match_text = ''.join(match_text)
                
                # Use fuzzy matching to identify duplicates
                if fuzz.ratio(search_text, str_match_text) > 95:
                    duplicate_info['index'].append(no)
                    break

            if len(duplicate_info['index']) > 1:
                break
            if i == len(doc.paragraphs) - 1:
                return duplicate_info
            
    # If a duplicate is found, extract its text
    if len(duplicate_info['index']) > 1:
        start = duplicate_info['index'][0]
        end = duplicate_info['index'][1] + 1
        text = ''.join([doc.paragraphs[row_no].text.strip().replace('\t', '') for row_no in range(start, end)])
        duplicate_info['text'].append(text)
    return duplicate_info



def remove_paragraphs(doc, indexes):
    """
    Remove paragraphs from the document based on the provided indexes.
    """
    # Sort indexes in reverse order to avoid shifting issues when removing paragraphs
    sorted_indexes = sorted(indexes, key=lambda x: x[0], reverse=True)
    for index in sorted_indexes:
        print(index)
        if len(index) > 1:
            for i in range(index[1], index[0]-1, -1):
                p = doc.paragraphs[i]._element
                p.getparent().remove(p)
                p._element = p._p = None



def remove_duplicate_paragraphs(doc, doc_type):
    """
    Remove duplicate paragraphs from the document.
    """
    #doc = Document(docx_path)

    # Find potential duplicate sections
    master_dict_text_index = find_duplicate_sections(doc, doc_type)
    print(f'master_dict_text_index {master_dict_text_index}')

    del_dup_index_list = []

    # For each potential duplicate section, search for actual duplicates
    for index, (start, end) in enumerate(master_dict_text_index['index']):
        current_start = end + 1
        while True:
            duplicate_info_dic = find_duplicate_text(doc, start_index=current_start, search_text=master_dict_text_index['text'][index])
            if len(duplicate_info_dic['index']) > 1:
                del_dup_index_list.append(duplicate_info_dic['index'])
                current_start = duplicate_info_dic['index'][-1] + 1
            else:
                break
    
    # Remove the identified duplicate paragraphs
    remove_paragraphs(doc, del_dup_index_list)
    #doc.save(docx_path.replace('.docx', '_1.docx'))





def new_get_pattern_index_of_text(start_index,doc):
        duplicate_info = {'text': [], 'index': []}
        secretary_of_state_index = []
        
        for i, para in enumerate(doc.paragraphs[start_index:], start=start_index):
            text = para.text.strip().lower()    
            
            if 'secretary of state' in text:
                secretary_of_state_index.append(i)
        
        for n in secretary_of_state_index:
            for i in range(n-1, -1, -1):
                text = doc.paragraphs[i].text.strip().lower()
                
                if 'title of rule' in text:
                    duplicate_info['index'].append([i, n-1])
                    break

            search_text = " ".join([p.text.strip().replace('\t', '') for p in doc.paragraphs[i:n]]).strip()
            duplicate_info['text'].append(search_text)

        return duplicate_info


def add_text_before_index(doc, para_indexes, text_to_insert):
    # Sort the indices in reverse order
    insert_positions = sorted([index[0] for index in para_indexes['index']], reverse=True)
    
    for i, position in enumerate(insert_positions):
        # Create a new paragraph before the current position
        new_paragraph = doc.paragraphs[position].insert_paragraph_before(text_to_insert)
        
        # Style the new paragraph
        run = new_paragraph.runs[0]  # Access the first run in the paragraph
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.bold = True
        new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add REDLINE at the beginning of the last entry
        if i == len(insert_positions) - 1:
            redline_run = new_paragraph.insert_paragraph_before().add_run('REDLINE')
            redline_font = redline_run.font
            redline_font.name = 'Times New Roman'
            redline_font.size = Pt(12)
            redline_font.bold = True
            

    return doc



def main(input_docx):
    print("entered redline text")
    doc = Document(input_docx)

    # Get the first header text
    # first_header_text = get_first_header_text(doc)
    # print("First header text:", first_header_text)

    # Clear all headers and footers
    # remove_headers_and_footers(doc)
    # print('removed header and footer')

    
    #*****The below functionality is for the title of rule and there vairants*****
    doc_type = ''
    index_of_title = 0
    
    # Find the index of "Title of Rule"
    for cntr, para in enumerate(doc.paragraphs):
        if 'title of rule:' in para.text.strip().lower():
            index_of_title = cntr
            break

    # Check the next few paragraphs for key phrases
    text_check = []
    for par in doc.paragraphs[index_of_title:index_of_title+10]:
        text_check.append(par.text)

    check_text_val = " ".join(text_check).strip().lower()

    title_and_rule_number = 'title of rule:' in check_text_val and 'rule number:' in check_text_val


    # Determine the document type based on key phrases
    if title_and_rule_number and 'division / contact / phone:' in check_text_val and 'secretary of state' in check_text_val:
        doc_type = 'TYPE 1'
    elif title_and_rule_number and 'division / contact / phone:' in check_text_val and 'statement of basis and purpose' in check_text_val:
        doc_type = 'TYPE 2'
    
    print(f'Type assigned {doc_type}')
    match doc_type:
        case 'TYPE 1' | 'TYPE 2':
            remove_duplicate_paragraphs(doc, doc_type)
        case 'TYPE 3':
            pass
    print("came till here")
                       

    dict_with_indexes = new_get_pattern_index_of_text(0, doc)

    # if dict_with_indexes['index'] and :
    #         doc = add_text_before_index(doc, dict_with_indexes, first_header_text)
    # else:
    #     print("add text call")
    #     add_text(doc,'')

    doc.save(input_docx)





# if __name__ == "__main__":
    
    # input_docx = r"C:\Users\6120867\Downloads\Netscan_Test_Files\co_files\NETSCAN_CO240701cas_Input\co_a001a207Redline.docx"
    # output_docx = input_docx.replace('.docx','_1.docx')
    # main(input_docx,output_docx)


