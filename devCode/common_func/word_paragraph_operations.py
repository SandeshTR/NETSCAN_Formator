from docx import Document
import logging
from docx.shared import Inches
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm , Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml.etree import QName
import re
import docx
from lxml import etree

def is_excluded_text(text):
    email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
    url_pattern = re.compile(r'\b(?:https?://|www\.)\S+\b')
    doc_link_pattern = re.compile(r'\b\S+\.(com|gov|aspx|html|txt|rtf|pdf|doc|docx|xls|xlsx)\b')
    return email_pattern.search(text) or url_pattern.search(text) or doc_link_pattern.search(text)

def word_file_indentation(doc_path):
    """ This function is used to set margins and selectively add first-line indentation to the word file """
    doc = Document(doc_path)
    
    # Loop through all sections and set margins to 1 inch
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Loop through all paragraphs
    for paragraph in doc.paragraphs:
        # Remove any existing indentation
        paragraph.paragraph_format.left_indent = Inches(0)
        paragraph.paragraph_format.right_indent = Inches(0)
        paragraph.paragraph_format.first_line_indent = Inches(0)
    
    doc.save(doc_path)
    logging.info(f"Margins adjusted and selective first-line indentation applied to the document: {doc_path}")

def format_paragraph(paragraph):
    """
    Apply specific formatting to a paragraph. Remove hyperlink text (but not images), apply text font and style.
    """
    replacements = {
        '–': '~#8211;',
        '§': '~#167;',
        ' ,': ',',
        '°': '~#176;',
        '•': '~#8226;',
        '✓': '/',
        '☒': '[x]',
        '→': '-->'
    }

    for run in paragraph.runs:
        # Check if the run contains an image
        if run._element.findall('.//w:drawing', namespaces=paragraph._element.nsmap):
            continue  # Skip this run if it contains an image

        if 'HYPERLINK' in run.element.xml:
            run.clear()  # Remove the hyperlink text

        if is_excluded_text(run.text):
            run.font.underline = False  # Remove underline if the text matches the exclusion criteria

        # replacing starting " with " and ending " with " and replacing all ' with ' 
        text = run.text
        text = re.sub(r'"(\b)', r'“\1', text)
        text = re.sub(r'(\b)"', r'”\1', text)
        text = re.sub(r'"([.,!?;:)\s]|$)', r'”\1', text)
        text = text.replace("'", "’")
        run.text = text

        # run.font.underline = False  # This will make all the underlines disappear
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def align_paragraph(paragraph):
    """
    Align the paragraph and set indentation.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.first_line_indent = Cm(0)

def apply_bullet_formatting(paragraph):
    """
    Apply bullet formatting if the paragraph is part of a bullet list.
    """
    if paragraph._element.xpath('.//w:numPr'):
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.first_line_indent = Cm(0)

def indent_table_left(table):
    """
    Indent the entire table to the left with 0 cm.
    """
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
    tblInd = tblPr.xpath('w:tblInd')
    if tblInd:
        tblPr.remove(tblInd[0])
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '0')
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)


def add_nac_to_section(doc, paragraph):
    """
    Add "NAC" to sections within the document.
    """
    section_pattern = re.compile(r'(Section|Sec\.?)\s*\d+\.\s*NAC\s*\d+[A-Z]?\.?\d+\s+is hereby amended to read as follows:')

    for i in range(len(doc.paragraphs) - 1):  # Iterate to the second-to-last paragraph
        paragraph = doc.paragraphs[i]
        if section_pattern.search(paragraph.text):
            next_paragraph = doc.paragraphs[i + 1]
            if not next_paragraph.text.strip().startswith("NAC"):
                next_paragraph.runs[0].text = "NAC " + next_paragraph.runs[0].text

def remove_spacing_between_paras(doc):
    """
    This function is used to remove spacing between paragraphs while preserving content
    """
    for para in doc.paragraphs:
        # Check if the paragraph is empty and doesn't contain any of these elements
        if (not para.text.strip()  and 
            not para._element.findall('.//w:drawing', namespaces=para._element.nsmap) and
            not para._element.findall('.//w:ins', namespaces=para._element.nsmap) and
            not para._element.findall('.//w:del', namespaces=para._element.nsmap) and
            not para._element.findall('.//w:smartTag', namespaces=para._element.nsmap)):
            
            if not para._element.getparent().tag.endswith('tbl'):
                p = para._element
                p.getparent().remove(p)
                p._element = p._p = None

def remove_section_breaks(doc_path):
    try:
        doc = docx.Document(doc_path)
        logging.info(f"Total paragraphs: {len(doc.paragraphs)}")

        for i, para in enumerate(doc.paragraphs[1:], 1):   
            text = para.text
            p_xml = etree.fromstring(para._p.xml)
            sect_prs = p_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
            #print(f"\nParagraph {i+1}:")

            del_elements = p_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}del')
            del_text_elements = p_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}delText')
            #w_namespace = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

            if not del_elements or not del_text_elements:
                if (text == "\x0c" or text == "" or text == chr(12) or 
                    "\x0c" in text or chr(12) in text or text.isspace()):
                    print("Text: "+text)
                    print(f"Found section break in paragraph {i+1}")
                    print(f"Text: '{text}'")
                    print(f"Length: {len(text)}")
                
                    p = para._element
                    for child in p[:]:  
                        if 'sectPr' in child.tag:
                            p.remove(child)
                    parent = p.getparent()
                    if parent is not None:
                        parent.remove(p)
                        print(f"Removed section break from paragraph")

            if sect_prs:
                for sect_pr in sect_prs:
                    parent = sect_pr.getparent()
                    if parent is not None:
                        parent.remove(sect_pr)
                
                new_p = etree.Element(para._p.tag, nsmap=para._p.nsmap)
                for element in p_xml:
                    new_p.append(element)
                if para._p.getparent() is not None:
                    para._p.getparent().replace(para._p, new_p)
                        #if new_text:
                    #para.text = new_text
                #else:
                    #para.text = ""
                
                #print(f"Removed section break from paragraph {i+1}")
        
        # Save the modified document
        doc.save(doc_path)
        logging.info(f"Successfully processed document")
        
    except Exception as e:
        logging.error(f"An error occurred: {str(e)}")
        raise e

# def remove_tabs_from_document(doc):
#     """Function to remove tabs from the file text while preserving images"""

#     def process_paragraph(paragraph):
#         for run in paragraph.runs:
#             if not run.element.findall('.//w:drawing', namespaces=run.element.nsmap):
#                 # Only process runs that don't contain images
#                 run.text = run.text.replace('\t', ' ')

#     # Process paragraphs
#     for paragraph in doc.paragraphs:
#         process_paragraph(paragraph)

#     # Process tables
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     process_paragraph(paragraph)

#     logging.info(f"Tabs removed successfully")