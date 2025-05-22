from docx import Document
import logging
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm , Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml.etree import QName
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import re
from docx.oxml.ns import nsmap
import win32com.client

def replace_image_with_text(doc_path, replacement_text,output_path):
    # Load the Word Document
    doc = Document(doc_path)
    
 
    wordprocessing_namespace = '{' + nsmap['wp'] + '}'

    # Iterate through all the paragraphs to find images
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
           
            for inline in run._element.findall('.//' + wordprocessing_namespace + 'inline'):
                
                inline.clear()
               
                run.text = replacement_text
    
   
    doc.save(output_path)
    logging.info('replaced image with tag non-displayable image text')

def remove_tables_and_insert_text(doc_path, replacement_text, output_path):
    document = Document(doc_path)

    
    max_threshold = 8.00
    value=0.00
    tables_to_replace = []

     
    for i, table in enumerate(document.tables):
        rows = len(table.rows)
        columns = len(table.columns)
        cells = rows * columns
        
        # Count merged cells (spans)
        spans = 0
        for row in table.rows:
            for cell in row.cells:
                # Count the number of merged cells
                if cell._element.xpath('.//w:vMerge'):
                    spans += 1
      

        value = (spans/cells)*100
        logging.info(f"Table has {rows} rows, {columns} columns, {cells} cells, and {spans} merged cells.{value} Percentage have cells")   
         
        if value>=max_threshold:        
                tables_to_replace.append((table))
                logging.info(f"Table index {i}")
                
    for  table in reversed(tables_to_replace):
        
        tbl_index = document.element.body.index(table._tbl)
        

        new_paragraph = document.add_paragraph(replacement_text)
        
        
        document.element.body.insert(tbl_index, new_paragraph._element)
        
        # Remove the table
        table._tbl.getparent().remove(table._tbl)

  
    document.save(output_path)

def remove_line_numbering(source_document,destination_document):
     
    word = win32com.client.Dispatch("Word.Application")
    doc = word.Documents.Open(source_document)
 
    # Iterate through sections to remove line numbering
    for section in doc.Sections:
         
        section.PageSetup.LineNumbering.Active = False
 
    # Save and close the document    
    doc.SaveAs(destination_document)
    doc.Close()
    word.Quit()
    logging.info('Removed line numbering')
 
def indent_table_left(table):
    """
    Indent the entire table to the left with 0 cm.
    """
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
    tblInd = tblPr.xpath('w:tblInd')
    if tblInd:
        tblPr.remove(tblInd[0])
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '0')
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

def limit_table_cell_length(doc_path,max_char=80):
    # Load the Word Document
    doc = Document(doc_path)
    
    # Iterate through each table in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells: 
                if len(cell.text) > max_char:
                        cell.text = cell.text[:max_char]
    
    # Save the modified document
    doc.save(doc_path)

def add_border(table):
    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                              f'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                              f'<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                              f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                              f'<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                              f'</w:tcBorders>')
        tcPr.append(tcBorders)
    logging.info('added border to table')

def add_border_doc(input_file,output_file):
  
    doc = Document(input_file)

    # Process each table in the document
    for table in doc.tables:
        # Add borders to the table
        add_border(table)
        
        # Set table alignment to center
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)  # Set font size
                #cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

   
    doc.save(output_file)