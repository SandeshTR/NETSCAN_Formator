import os
import fitz  # PyMuPDF
import docx
import pandas as pd
from collections import defaultdict
import re
import difflib

class FormatComparator:
    def __init__(self, pdf_path, docx_path):
        """Initialize with paths to PDF and DOCX files"""
        self.pdf_path = pdf_path
        self.docx_path = docx_path
        self.issues = []
        
    def extract_from_pdf(self):
        """Extract text with formatting from PDF"""
        doc = fitz.open(self.pdf_path)
        content = []
        
        for page_num, page in enumerate(doc):
            # Extract text with formatting info
            text_dict = page.get_text("dict")
            
            # Process each text block
            for block in text_dict["blocks"]:
                if block["type"] == 0:  # Text block
                    for line in block["lines"]:
                        for span in line["spans"]:
                            # Check formatting flags
                            is_bold = span["font"].lower().find("bold") >= 0 or (span["flags"] & 2**4 != 0)
                            is_italic = span["font"].lower().find("italic") >= 0 or (span["flags"] & 2**1 != 0)
                            is_underlined = span["flags"] & 2**2 != 0
                            is_strikethrough = span["flags"] & 2**3 != 0
                            
                            content.append({
                                "text": span["text"],
                                "page": page_num + 1,
                                "bold": is_bold,
                                "italic": is_italic,
                                "underline": is_underlined,
                                "strikethrough": is_strikethrough,
                                "origin": (span["origin"][0], span["origin"][1]),
                                "bbox": span["bbox"]
                            })
        
        return content
    
    def extract_from_docx(self):
        """Extract text with formatting from DOCX"""
        doc = docx.Document(self.docx_path)
        content = []
        
        # Process paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
                
            # Process runs (segments with consistent formatting)
            for run in para.runs:
                if not run.text.strip():
                    continue
                    
                content.append({
                    "text": run.text,
                    "paragraph": para_idx,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "strikethrough": run.font.strike,
                })
        
        # Process tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        if not para.text.strip():
                            continue
                            
                        for run in para.runs:
                            if not run.text.strip():
                                continue
                                
                            content.append({
                                "text": run.text,
                                "table": table_idx,
                                "row": row_idx,
                                "cell": cell_idx,
                                "paragraph": para_idx,
                                "bold": run.bold,
                                "italic": run.italic,
                                "underline": run.underline,
                                "strikethrough": run.font.strike,
                            })
        
        return content
    
    def normalize_text(self, text):
        """Normalize text for comparison by removing extra whitespace"""
        return re.sub(r'\s+', ' ', text).strip()
    
    def find_matching_text(self, text, content_list, threshold=0.8):
        """Find matching text in content list using fuzzy matching"""
        normalized_text = self.normalize_text(text)
        best_match = None
        best_score = threshold
        
        for item in content_list:
            normalized_item_text = self.normalize_text(item["text"])
            if normalized_text == normalized_item_text:
                return item, 1.0
            
            # Use difflib for fuzzy matching
            score = difflib.SequenceMatcher(None, normalized_text, normalized_item_text).ratio()
            if score > best_score:
                best_score = score
                best_match = item
        
        return best_match, best_score
    
    def compare_formatting(self):
        """Compare formatting between PDF and DOCX"""
        pdf_content = self.extract_from_pdf()
        docx_content = self.extract_from_docx()
        
        # Group PDF content by page for reporting
        pdf_by_page = defaultdict(list)
        for item in pdf_content:
            pdf_by_page[item["page"]].append(item)
        
        # Compare each PDF text segment with DOCX
        for pdf_item in pdf_content:
            # Skip very short text segments (often punctuation or spaces)
            if len(pdf_item["text"].strip()) < 3:
                continue
                
            # Find matching text in DOCX
            docx_match, score = self.find_matching_text(pdf_item["text"], docx_content)
            
            if docx_match and score > 0.8:  # Good match found
                # Check for formatting differences
                for format_type in ["bold", "italic", "underline", "strikethrough"]:
                    pdf_format = pdf_item.get(format_type, False)
                    docx_format = docx_match.get(format_type, False)
                    
                    # If formatting differs, record the issue
                    if pdf_format != docx_format and pdf_format:  # Only report if PDF has formatting that DOCX doesn't
                        self.issues.append({
                            "issue_type": f"Missing {format_type}",
                            "pdf_page": pdf_item["page"],
                            "pdf_text": pdf_item["text"],
                            "docx_text": docx_match["text"],
                            "description": f"Text in PDF has {format_type} formatting that is missing in DOCX"
                        })
        
        return self.issues
    
    def generate_report(self):
        """Generate a report of formatting issues"""
        if not self.issues:
            self.compare_formatting()
            
        if not self.issues:
            print("No formatting issues found!")
            return pd.DataFrame()
            
        # Convert issues to DataFrame for easy viewing/export
        df = pd.DataFrame(self.issues)
        
        # Group by page and issue type
        summary = df.groupby(['pdf_page', 'issue_type']).size().reset_index(name='count')
        
        print(f"Found {len(self.issues)} formatting issues!")
        print("\nSummary by page and issue type:")
        print(summary)
        
        return df
    
    def save_report(self, output_path="formatting_issues.xlsx"):
        """Save the report to Excel file"""
        df = self.generate_report()
        if not df.empty:
            df.to_excel(output_path, index=False)
            print(f"Report saved to {output_path}")


if __name__ == "__main__":
    # Replace with your file paths
    pdf_path = r"C:\File\NETSCAN\Input\co_p001aft183.pdf"
    docx_path = r"C:\File\NETSCAN\Input\co_p001aft183.docx"
    
    comparator = FormatComparator(pdf_path, docx_path)
    issues_df = comparator.generate_report()
    comparator.save_report()