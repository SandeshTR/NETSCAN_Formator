import docx
import re
from docx.shared import RGBColor
import logging

def has_font_effects(paragraph):

    for run in paragraph.runs:
        if run.font.small_caps:
            return True
    return False

def is_all_caps(text):
    return text.isupper()

def has_word_effects(word):
    return word.font.small_caps

def apply_small_caps(word):
    word.font.small_caps = True

def is_version_or_section_number(text):
    return bool(re.match(r'^\s*[\d.]+\s*$', text))

def process_paragraph(paragraph):
    if has_font_effects(paragraph):
        logging.info("Paragraph: " + paragraph.text)
        
        
        has_small_caps = any(run.font.small_caps for run in paragraph.runs)
        has_non_effect = any(not run.font.small_caps for run in paragraph.runs)
        
        if has_small_caps and has_non_effect:
        
            non_effect_words = [run.text for run in paragraph.runs if not run.font.small_caps]
            #non_effect_words = [
            #    word for run in paragraph.runs 
            #    if not run.font.small_caps
            #    for word in re.findall(r'\b[a-zA-Z]+\b', run.text)
            #]
            
            all_non_effect_upper = all(word.isupper() for word in non_effect_words if word.strip() and not re.search(r'\d', word))
            logging.info(non_effect_words)
            logging.info(all_non_effect_upper)

            if all_non_effect_upper:
                logging.info("Converting non-effect uppercase words to small caps")
                for run in paragraph.runs:
                    if not run.font.small_caps:
                        if run.text.isupper():
                            run.text = run.text.lower()
                            apply_small_caps(run)

def check_small_caps(doc_path):
    doc=docx.Document(doc_path)
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    doc.save(doc_path)


# Usage
#doc_path = r"C:\File\NETSCAN\Input\co_p007aft197.docx"
#processed_filename = r"C:\File\NETSCAN\Input\co_p007aft197_processed.docx"
#doc = docx.Document(filename)
#check_small_caps(doc)
#process_doc(doc)
#doc.save(processed_filename)