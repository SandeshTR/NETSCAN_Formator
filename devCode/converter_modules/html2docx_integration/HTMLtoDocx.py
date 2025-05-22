from bs4 import BeautifulSoup
from docx import Document
import os
import re
from logs.logs_handler import get_logger

logger = get_logger(__name__)

def HTMLtoWord(html_file_path, word_file_path):
    # Determine document type from filename
    doc_type = get_document_type(html_file_path)
    
    # Process the HTML to extract relevant content
    processed_html_path = html_file_path.replace(".html", "_processed.html")
    HTMLParsing(html_file_path, processed_html_path, doc_type)
    
    # Read the processed HTML content
    try:
        with open(processed_html_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
    except Exception as e:
        logger.warning(f"Error reading processed HTML file: {e}")
        return

    # Parse the HTML content with BeautifulSoup
    soup = BeautifulSoup(html_content, 'lxml')
    
    # Create a new Word document
    doc = Document()
    
    # Process all elements and add them to the document
    process_elements(soup.body, doc)
    
    # Save the Word document
    try:
        doc.save(word_file_path)
        logger.info(f'HTML content has been converted and saved to {word_file_path}')
    except Exception as e:
        logger.warning(f"Error saving Word document: {e}")
    
    # Clean up the temporary file
    try:
        os.remove(processed_html_path)
    except Exception as e:
        logger.warning(f"Warning: Could not remove temporary file: {e}")

def get_document_type(file_path):
    """Determine document type from filename"""
    if "_co_p" in file_path:
        return "proposed"
    elif "_co_a" in file_path:
        return "adopted"
    elif "_co_e" in file_path:
        return "emergency"
    elif "_coac" in file_path:
        return "admin_change"
    elif "_co_c" in file_path:
        return "correction"
    else:
        return "standard"

def process_elements(parent_element, doc):
    """Process all elements and add them to the Word document"""
    if not parent_element:
        return
        
    for element in parent_element.find_all(True, recursive=False):
        if element.name == 'h1' or element.name == 'h2':
            doc.add_heading(element.get_text(strip=True), level=1)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(strip=True), level=2)
        elif element.name == 'p':
            text = element.get_text(strip=True)
            if text:  # Only add if there's actual text
                para = doc.add_paragraph()
                run = para.add_run(text)
                if element.find(['strong', 'b']):
                    run.bold = True
        elif element.name == 'hr':
            doc.add_paragraph('---')
        elif element.name == 'table':
            process_table(element, doc)
        elif element.name == 'div':
            # Process div contents
            process_elements(element, doc)
        elif element.name == 'ul' or element.name == 'ol':
            process_list(element, doc)
        elif element.name == 'img':
            img_url = element.get('src', '')
            if img_url:
                doc.add_paragraph(f'[Image: {img_url}]')
        # Process any children of this element
        if element.name not in ['table', 'ul', 'ol']:  # These are handled separately
            process_elements(element, doc)

def process_table(table_element, doc):
    """Process a table element and add it to the Word document"""
    rows = table_element.find_all('tr')
    if not rows:
        return
        
    # Count max columns
    max_cols = 0
    for row in rows:
        cols = len(row.find_all(['td', 'th']))
        max_cols = max(max_cols, cols)
    
    if max_cols == 0:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    
    # Track which rows to skip (will be marked True if should be skipped)
    skip_rows = [False] * len(rows)
    
    # First pass: identify rows to skip based on business rules
    for row_idx, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        
        # Check for cells with text that should trigger row removal
        for cell in cells:
            cell_text = cell.get_text(strip=True).lower()
            
            # Check for "Rule" or "Adopted Rules" rows that should be removed
            if cell_text == "rule" or cell_text == "adopted rules":
                skip_rows[row_idx] = True
                break
    
    # Second pass: process rows that shouldn't be skipped
    row_offset = 0  # To track the shifting row indices due to skipped rows
    
    for row_idx, row in enumerate(rows):
        if skip_rows[row_idx]:
            row_offset += 1
            continue
            
        cells = row.find_all(['td', 'th'])
        
        for col_idx, cell in enumerate(cells):
            if col_idx >= max_cols:
                continue
                
            cell_text = cell.get_text(strip=True)
            
            # Apply business rules for specific cell text
            cell_text = apply_business_rules_to_cell(cell_text)
                
            # Add text to the cell
            table_cell = table.cell(row_idx - row_offset, col_idx)
            paragraph = table_cell.paragraphs[0]
            run = paragraph.add_run(cell_text)
            
            # Apply formatting
            if cell.find(['strong', 'b']):
                run.bold = True
            if cell.find(['i', 'em']):
                run.italic = True

def apply_business_rules_to_cell(cell_text):
    """Apply business rules to specific cell text"""
    # Check for text that needs "(see below if available)" appended
    special_texts = [
        "additional information",
        "basis and purpose",
        "purpose/objective of rule basis and purpose",
        "redline",
        "emergency justification"
    ]
    
    cell_lower = cell_text.lower()
    
    for special_text in special_texts:
        if cell_lower == special_text:
            return cell_text + " (see below if available)"
    
    return cell_text

def process_list(list_element, doc):
    """Process a list element and add it to the Word document"""
    for idx, item in enumerate(list_element.find_all('li', recursive=False)):
        prefix = "• " if list_element.name == 'ul' else f"{idx+1}. "
        para = doc.add_paragraph()
        para.add_run(prefix + item.get_text(strip=True))

def HTMLParsing(html_file_path, output_html_file_path, doc_type):
    """Extract and process relevant content from the HTML file"""
    try:
        with open(html_file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
    except Exception as e:
        logger.warning(f"Error reading HTML file: {e}")
        return
    
    # Parse the HTML content
    soup = BeautifulSoup(html_content, 'lxml')
    
    # Create output HTML structure
    output_soup = BeautifulSoup('<html><head><title>Processed Content</title></head><body></body></html>', 'lxml')
    output_body = output_soup.body
    
    # Extract tracking number details if present
    tracking_details = extract_tracking_details(soup)
    if tracking_details:
        heading = output_soup.new_tag('h2')
        heading.string = tracking_details
        output_body.append(heading)
    
    # Find all section headers (darkBlueText spans)
    sections = soup.find_all('span', class_='darkBlueText')
    
    # If no sections found with darkBlueText, try to find tables directly
    if not sections:
        # Find content tables
        tables = soup.find_all('table')
        content_tables = []
        
        for i, table in enumerate(tables):
            if is_content_table(table):
                # Apply business rules to the table before adding it
                apply_business_rules_to_table(table, doc_type)
                content_tables.append(table)
        
        # Add content tables to output
        for table in content_tables:
            output_body.append(table)
    else:
        # Process each section with its associated table
        for section in sections:
            section_name = section.get_text(strip=True)
            
            # Create heading for this section
            heading = output_soup.new_tag('p')
            heading.string = section_name
            output_body.append(heading)
            
            # Find the table that follows this section
            next_table = find_next_content_table(section)
            if next_table:
                # Apply business rules to the table before adding it
                apply_business_rules_to_table(next_table, doc_type)
                output_body.append(next_table)
    
    # Write the output HTML
    try:
        with open(output_html_file_path, 'w', encoding='utf-8') as file:
            file.write(str(output_soup))
    except Exception as e:
        logger.warning(f"Error writing processed HTML file: {e}")

def apply_business_rules_to_table(table, doc_type):
    """Apply business rules to a table based on document type"""
    # Find rows to remove based on business rules
    rows_to_remove = []
    
    for row in table.find_all('tr'):
        cells = row.find_all(['td', 'th'])
        if len(cells) >= 2:
            first_cell_text = cells[0].get_text(strip=True).lower()
            
            # For proposed rule documents: Delete "Rule" line
            if doc_type == "proposed" and first_cell_text == "rule":
                rows_to_remove.append(row)
            
            # For adopted rule documents: Delete "Rule" and "Adopted Rules" lines
            elif doc_type == "adopted" and (first_cell_text == "rule" or first_cell_text == "adopted rules"):
                rows_to_remove.append(row)
    
    # Remove the identified rows
    for row in rows_to_remove:
        row.decompose()
    
    # Modify cells based on business rules
    for row in table.find_all('tr'):
        cells = row.find_all(['td', 'th'])
        if len(cells) >= 2:
            first_cell_text = cells[0].get_text(strip=True).lower()
            
            # Add "(see below if available)" to specific cells
            if first_cell_text in ["additional information", "basis and purpose", 
                                  "purpose/objective of rule basis and purpose", 
                                  "redline", "emergency justification"]:
                # Check if text already contains the suffix
                if "(see below if available)" not in cells[0].get_text():
                    cells[0].string = cells[0].get_text(strip=True) + " (see below if available)"

def extract_tracking_details(soup):
    """Extract tracking number details from the document"""
    # Try to find tracking number details in text
    for cell in soup.find_all('td'):
        cell_text = cell.get_text()
        if 'Details of Tracking Number' in cell_text:
            # Extract text between "Details of Tracking Number" and "CCR details"
            start_marker = 'Details of Tracking Number'
            end_marker = 'CCR details'
            
            start_idx = cell_text.find(start_marker)
            end_idx = cell_text.find(end_marker, start_idx)
            
            if start_idx != -1 and end_idx != -1:
                return cell_text[start_idx:end_idx].strip().replace('\n\t\xa0', '')
    
    # Alternative method: look for a specific heading
    heading = soup.find('h2', string=lambda s: s and 'Details of Tracking Number' in s)
    if heading:
        return heading.get_text(strip=True)
    
    return None

def find_next_content_table(element):
    """Find the next content table after an element"""
    current = element
    while current:
        current = current.find_next()
        if current and current.name == 'table':
            return current
    return None

def is_content_table(table):
    """Determine if a table contains actual content rather than navigation"""
    # Tables with many rows are likely content tables
    if len(table.find_all('tr')) > 3:
        return True
        
    # Tables with cells containing substantial text are likely content tables
    for cell in table.find_all(['td', 'th']):
        text = cell.get_text(strip=True)
        if len(text) > 50:  # Arbitrary threshold for "substantial" text
            return True
            
    # Check for content-related keywords
    keywords = ['ccr', 'rule', 'section', 'adopted', 'proposed', 'details', 
                'rule details', 'hearing information', 'contact information']
    table_text = table.get_text().lower()
    if any(keyword in table_text for keyword in keywords):
        return True
        
    return False

# Example usage
#if __name__ == "__main__":
#    input_path = r"C:\File\NETSCAN\Input\HTML\24-E08.html"
#    output_path = input_path.replace(".html", ".docx")
#    HTMLtoWord(input_path, output_path)