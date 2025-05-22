from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import RGBColor
import chardet
from logs.logs_handler import get_logger

logger = get_logger(__name__)

def detect_encoding(file_path):
    """
    Detects the encoding of a file using chardet.
    Returns the detected encoding or 'utf-8' if detection fails.
    """
    try:
        with open(file_path, 'rb') as f:
            raw_data = f.read()
        result = chardet.detect(raw_data)
        encoding = result['encoding']
        confidence = result['confidence']
        if encoding and confidence > 0.5:
            logger.info(f"Detected encoding for {file_path}: {encoding} (Confidence: {confidence})")
            return encoding
        else:
            logger.warning(f"Low confidence in encoding detection for {file_path}. Falling back to 'utf-8'.")
            return 'utf-8'
    except Exception as e:
        logger.error(f"Error detecting encoding for {file_path}: {e}. Falling back to 'utf-8'.")
        return 'utf-8'

def HTMLParsing(html_path, output_html_file_path):
    """
    Parses the HTML file, extracts relevant tables or elements,
    and writes to a temporary HTML file.
    """
    try:
        encoding = detect_encoding(html_path)
        
        with open(html_path, 'r', encoding=encoding, errors='replace') as file:
            html_content = file.read()
        
        soup = BeautifulSoup(html_content, 'lxml')
        tables = soup.find_all('table')
        
        output_soup = BeautifulSoup('<html><head><title>Extracted Content</title></head><body></body></html>', 'lxml')
        output_body = output_soup.body
        extracted_texts = []
        
        # Initialize flag and index to check for 'Details of Tracking Number'
        has_tracking_number = False
        tracking_number_table_index = None
        
        # Step 1: Check if the third table contains 'Details of Tracking Number'
        if len(tables) >= 3:
            third_table = tables[2]
            cells = third_table.find_all('td')
            for cell in cells:
                if 'Details of Tracking Number' in cell.get_text():
                    has_tracking_number = True
                    tracking_number_table_index = 2
                    logger.info(f"'Details of Tracking Number' found in table index 2 of {html_path}")
                    break
        
        # Step 2: If not found in the third table, search the entire HTML
        if not has_tracking_number:
            # Search all tables for 'Details of Tracking Number'
            for idx, table in enumerate(tables):
                cells = table.find_all('td')
                for cell in cells:
                    if 'Details of Tracking Number' in cell.get_text():
                        has_tracking_number = True
                        tracking_number_table_index = idx
                        logger.info(f"'Details of Tracking Number' found in table index {idx} of {html_path}")
                        break
                if has_tracking_number:
                    break
        
        if has_tracking_number:
            # Step 3: Process tables based on whether 'Details of Tracking Number' was found
            for table_index, table in enumerate(tables):
                if has_tracking_number:
                    # Skip tables before the tracking number table
                    if table_index < tracking_number_table_index:
                        continue  # Skip tables before the tracking number table
                    
                    if table_index == tracking_number_table_index:
                        rows = table.find_all('tr')
                        for row in rows:
                            cells = row.find_all('td')
                            for cell in cells:
                                if 'Details of Tracking Number' in cell.get_text():
                                    data = cell.get_text()
                                    start_marker = 'Details of Tracking Number'
                                    end_marker = 'CCR details'
                                    start_index = data.find(start_marker)
                                    end_index = data.find(end_marker, start_index)
                                    
                                    soup1 = BeautifulSoup(str(table), 'html.parser')
                                    span_tags = soup1.find_all('span', class_='darkBlueText')
                                    
                                    for span in span_tags:
                                        text = span.get_text(strip=True)
                                        if text:  # Check if text is not empty
                                            extracted_texts.append(text)
                                    
                                    if extracted_texts:
                                        logger.info(f"Extracted texts: {extracted_texts[0]}")
                                    
                                    if start_index != -1 and end_index != -1:
                                        extracted_text = data[start_index:end_index].strip()
                                        extracted_text = ' '.join(extracted_text.split())  # Remove excess whitespace
                                        if extracted_text:  # Only add if not empty
                                            heading1 = output_soup.new_tag('p')
                                            heading1.string = extracted_text.replace('\n\t\xa0', '')
                                            output_body.append(heading1)
                        continue  # Move to next table after processing
                    
                    elif table_index > tracking_number_table_index:
                        heading_text = ""
                        if (table_index - tracking_number_table_index -1) < len(extracted_texts):
                            heading_text = extracted_texts[table_index - tracking_number_table_index -1]
                        else:
                            heading_text = f"Table {table_index + 1}"
                        
                        heading_text = ' '.join(heading_text.split())  # Remove excess whitespace
                        if heading_text:  # Only add if not empty
                            heading = output_soup.new_tag('p')
                            heading.string = heading_text
                            output_body.append(heading)
                        
                        if table:
                            output_body.append(table)
        else:
            # **New Logic: When 'Details of Tracking Number' is not found**
            logger.info(f"'Details of Tracking Number' not found in any table of {html_path}. Processing all relevant elements.")
            body = soup.body
            if not body:
                logger.warning(f"No <body> tag found in {html_path}. Skipping.")
            else:
                for element in body.find_all(['h2', 'h3', 'p', 'table', 'br']):
                    if element.name in ['h2', 'h3', 'p']:
                        # For <h2> elements, remove line breaks before appending
                        if element.name == 'h2':
                            original_text = element.get_text(separator=' ', strip=True)
                            cleaned_text = ' '.join(original_text.split())
                            new_tag = output_soup.new_tag(element.name)
                            new_tag.string = cleaned_text
                            output_body.append(new_tag)
                            logger.info(f"Processed <h2> element: {cleaned_text}")
                        else:
                            # Append <h3> and <p> elements as-is
                            output_body.append(element)
                    elif element.name == 'table':
                        # Append tables to the output_body
                        output_body.append(element)
                    elif element.name == 'br':
                        # Handle line breaks if necessary
                        br_tag = output_soup.new_tag('br')
                        output_body.append(br_tag)
    
        with open(output_html_file_path, 'w', encoding='utf-8') as file:
            file.write(str(output_soup.prettify()))
        
        logger.info(f"Extracted content written to temporary HTML file: {output_html_file_path}")
    except Exception as e:
        logger.error(f"Error during HTML parsing for {html_path}: {e}")
        raise  # Re-raise exception to handle it in calling function


def add_element_to_doc(element, doc, html_file_path):
    """
    Adds HTML elements to the Word document based on their tag.
    All text is set to black color. Skips specific table rows based on file name and cell text.
    """
    if element.name == 'h1':
        text = element.get_text(strip=True)
        if text:
            heading = doc.add_heading(text, level=1)
            # Set font color to black
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
    elif element.name == 'h2':
        text = element.get_text(separator=' ', strip=True)
        cleaned_text = ' '.join(text.split())  # Remove excess whitespace and line breaks
        if cleaned_text:
            heading = doc.add_heading(cleaned_text, level=2)
            # Set font color to black
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
    elif element.name == 'h3':
        text = element.get_text(strip=True)
        if text:
            heading = doc.add_heading(text, level=3)
            # Set font color to black
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
    elif element.name == 'p':
        text = element.get_text(strip=True)
        if text:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.bold = True
            # Set font color to black
            run.font.color.rgb = RGBColor(0, 0, 0)
    elif element.name == 'hr':
        # Only add horizontal line representation if necessary
        paragraph = doc.add_paragraph('---')
        run = paragraph.runs[0]
        # Set font color to black
        run.font.color.rgb = RGBColor(0, 0, 0)
    elif element.name == 'table':
        try:
            rows = element.find_all('tr')
            if not rows:
                return

            filename = html_file_path.name.lower()

            # Filter out rows that should be skipped
            included_rows = []
            for row in rows:
                cells = row.find_all(['td', 'th'])
                skip_row = False
                for cell in cells:
                    cell_text = cell.get_text(separator=' ', strip=True).lower()
                    cell_text = ' '.join(cell_text.split())
                    if (('co_p' in filename) and (cell_text == 'rule')) or \
                       (('co_a' in filename) and (cell_text in ['rule', 'adopted rules'])) or \
                       (('co_e' in filename) and (cell_text == 'adopted rules')):
                        skip_row = True                        
                        break
                if not skip_row:
                    included_rows.append(row)

            if not included_rows:
                return  # No rows to include

            # Determine the maximum number of columns
            cols = max(len(row.find_all(['td', 'th'])) for row in included_rows)

            # Create the table with the included rows
            table_doc = doc.add_table(rows=len(included_rows), cols=cols)
            table_doc.style = 'Table Grid'

            for row_idx, row in enumerate(included_rows):
                cells = row.find_all(['td', 'th'])
                for col_idx in range(cols):
                    if col_idx < len(cells):
                        cell = cells[col_idx]
                        cell_text = cell.get_text(strip=True)
                        if cell_text:  # Only add text if not empty
                            cell_paragraph = table_doc.cell(row_idx, col_idx).paragraphs[0]
                            run = cell_paragraph.add_run(cell_text)
                            
                            if cell.find(['strong', 'b']):
                                run.bold = True
                            if cell.find(['i', 'em']):
                                run.italic = True
                            
                            # Set font color to black
                            run.font.color.rgb = RGBColor(0, 0, 0)
        except Exception as e:
            logger.error(f"Error processing table in Word document: {e}")
    elif element.name == 'img':
        img_url = element.get('src')
        if img_url:
            paragraph = doc.add_paragraph(f'Image: {img_url}')
            run = paragraph.runs[0]
            # Set font color to black
            run.font.color.rgb = RGBColor(0, 0, 0)


def HTMLtoWord(html_file_path, word_file_path):
    """
    Converts a single HTML file to a Word document.
    """
    try:

        html_file_path = Path(html_file_path) if not isinstance(html_file_path, Path) else html_file_path
        word_file_path = Path(word_file_path) if not isinstance(word_file_path, Path) else word_file_path
        # Define the temporary HTML file path
        temp_html_file_path = str(html_file_path).replace(".html","_1.html")
        
        # Parse and extract necessary parts of the HTML
        HTMLParsing(html_file_path, temp_html_file_path)
        
        # Detect encoding for the temporary HTML file
        encoding = detect_encoding(temp_html_file_path)
        
        # Read the modified HTML content from the temporary file
        with open(temp_html_file_path, 'r', encoding=encoding, errors='replace') as file:
            html_content = file.read()
        
        # Parse the HTML content with BeautifulSoup
        soup = BeautifulSoup(html_content, 'lxml')
        
        # Create a new Word document
        doc = Document()
        
        # Iterate over all elements in the HTML body and add them to the Word document
        for element in soup.body.find_all(True):  # Find all tags within body
            add_element_to_doc(element, doc,html_file_path)
        
        # Save the Word document
        doc.save(word_file_path)
        logger.info(f"Converted: {html_file_path} -> {word_file_path}")
        
    except Exception as e:
        logger.error(f"Failed to convert {html_file_path}: {e}")
    finally:
        # Attempt to remove the temporary HTML file
        try:
            if Path(temp_html_file_path).exists():
                Path(temp_html_file_path).unlink()
                logger.info(f"Removed temporary file: {temp_html_file_path}")
        except Exception as e:
            logger.warning(f"Could not remove temporary file {temp_html_file_path}: {e}")

#if __name__ == "__main__":
    # Define your input, output, and log folders
    #html_file = r"C:\File\NETSCAN\Input\HTML\2025-00034_coac.html"    # Replace with your actual input folder path
    #word_file = r"C:\File\NETSCAN\Input\HTML\2025-00034_coac.docx"  # Replace with your actual output folder path
    
    #HTMLtoWord(html_file, word_file)

    #process_all_html_files(input_directory, output_directory, log_directory)